"""
Page Configuration API Routes
Public endpoints for live config and Admin endpoints for CRUD operations.
"""
from flask import request, Response, g
from flask_jwt_extended import jwt_required, current_user

from app.api.page_config import page_config_bp
from app.api.page_config.service import (
    PageConfigService, AssetService, get_default_config, apply_translations
)
from app.services.s3_service import S3Service
from app.common.decorators import feature_required, role_required
from app.common.responses import (
    success_response, error_response, not_found_response, created_response
)
from app.models import PageType, UserRole, Tenant


# ============== PUBLIC ENDPOINTS ==============

@page_config_bp.route('/public/<page_type>', methods=['GET'])
def get_public_config(page_type):
    """
    Get the LIVE page configuration for public display.
    No authentication required.

    Query params:
        lang (str): Language code ('en', 'hi', 'te'). Defaults to 'en'.

    Returns default config if no LIVE config exists.
    Only serves languages listed in published_languages.
    """
    # Validate page type
    try:
        PageType(page_type)
    except ValueError:
        return error_response(f"Invalid page type: {page_type}. Valid types: {[p.value for p in PageType]}")

    lang = request.args.get('lang', 'en')

    # Resolve which tenant's LIVE config to serve.
    #
    # The tenant is normally resolved SERVER-SIDE from the request host by the
    # ``before_request`` hook (``g.tenant_id`` / ``g.tenant_source``). A
    # client-supplied ``?tenant_slug=`` may only *fill in* a tenant when the
    # host did NOT already resolve one (``tenant_source == 'default_fallback'``
    # — e.g. an explicit preview from the apex/localhost). It must NEVER
    # override a host-resolved tenant.
    #
    # SECURITY: the frontend's old slug resolver returned the apex default
    # ``'platform'`` for every ``www.<tenant>`` host, and honouring that here
    # unconditionally pulled the APEX login config onto real tenants' domains
    # (cross-tenant leak). Gating on ``tenant_source`` closes that: a real
    # host-resolved tenant now always wins over the client hint.
    tenant_slug = request.args.get('tenant_slug')
    if tenant_slug and getattr(g, 'tenant_source', None) == 'default_fallback':
        tenant = Tenant.query.filter_by(slug=tenant_slug).first()
        if not tenant:
            return not_found_response(f'Tenant "{tenant_slug}"')
        g.tenant_id = tenant.id  # type: ignore[attr-defined]

    config = PageConfigService.get_live_config(page_type)

    if config:
        config_dict = config.to_dict(include_asset_urls=True)

        # Check if requested language is published
        published = config_dict.get('published_languages') or ['en']
        if lang != 'en' and lang not in published:
            # Fallback to English if requested language not published
            lang = 'en'

        # Apply translations for the requested language
        config_dict = apply_translations(config_dict, lang)
        config_dict['current_language'] = lang
        config_dict['available_languages'] = published

        return success_response(config_dict)

    # Return default config
    default = get_default_config(page_type)
    default['current_language'] = 'en'
    default['available_languages'] = ['en']
    return success_response(default)


@page_config_bp.route('/public/types', methods=['GET'])
def get_page_types():
    """Get list of available page types."""
    return success_response([p.value for p in PageType])


# ============== ADMIN ENDPOINTS ==============

@page_config_bp.route('/admin/<page_type>', methods=['GET'])
@jwt_required()
@role_required(UserRole.SUPER_ADMIN)
@feature_required('admin.page_configuration')
def admin_get_config(page_type):
    """
    Get all configs (draft, preview, live) for a page type.
    Super Admin only.
    """
    try:
        PageType(page_type)
    except ValueError:
        return error_response(f"Invalid page type: {page_type}")
    
    draft = PageConfigService.get_draft_config(page_type)
    preview = PageConfigService.get_preview_config(page_type)
    live = PageConfigService.get_live_config(page_type)
    
    return success_response({
        'draft': draft.to_dict(include_asset_urls=True) if draft else None,
        'preview': preview.to_dict(include_asset_urls=True) if preview else None,
        'live': live.to_dict(include_asset_urls=True) if live else None,
    })


@page_config_bp.route('/admin/<page_type>/draft', methods=['GET'])
@jwt_required()
@role_required(UserRole.SUPER_ADMIN)
@feature_required('admin.page_configuration')
def get_draft(page_type):
    """Get or create draft configuration."""
    try:
        user_id = str(current_user.id) if current_user else None
        draft = PageConfigService.get_or_create_draft(page_type, user_id)
        return success_response(draft.to_dict(include_asset_urls=True))
    except ValueError as e:
        return error_response(str(e))


@page_config_bp.route('/admin/<page_type>/draft', methods=['PUT'])
@jwt_required()
@role_required(UserRole.SUPER_ADMIN)
@feature_required('admin.page_configuration')
def update_draft(page_type):
    """
    Update draft configuration.
    
    Request body can contain any of the configurable fields.
    """
    data = request.get_json()
    if not data:
        return error_response("Request body required")
    
    try:
        user_id = str(current_user.id) if current_user else None
        draft = PageConfigService.update_draft(page_type, data, user_id)
        return success_response(
            draft.to_dict(include_asset_urls=True),
            message="Draft updated successfully"
        )
    except ValueError as e:
        return error_response(str(e))


@page_config_bp.route('/admin/<page_type>/preview', methods=['POST'])
@jwt_required()
@role_required(UserRole.SUPER_ADMIN)
@feature_required('admin.page_configuration')
def promote_to_preview(page_type):
    """Promote draft to preview status."""
    try:
        user_id = str(current_user.id) if current_user else None
        preview = PageConfigService.promote_to_preview(page_type, user_id)
        return success_response(
            preview.to_dict(include_asset_urls=True),
            message="Draft promoted to preview"
        )
    except ValueError as e:
        return error_response(str(e))


@page_config_bp.route('/admin/<page_type>/preview', methods=['GET'])
@jwt_required()
@role_required(UserRole.SUPER_ADMIN)
@feature_required('admin.page_configuration')
def get_preview(page_type):
    """Get preview configuration."""
    preview = PageConfigService.get_preview_config(page_type)
    if not preview:
        return not_found_response(f"Preview config for {page_type}")
    return success_response(preview.to_dict(include_asset_urls=True))


@page_config_bp.route('/admin/<page_type>/publish', methods=['POST'])
@jwt_required()
@role_required(UserRole.SUPER_ADMIN)
@feature_required('admin.page_configuration')
def publish_config(page_type):
    """Publish preview to live."""
    try:
        user_id = str(current_user.id) if current_user else None
        live = PageConfigService.publish(page_type, user_id)
        return success_response(
            live.to_dict(include_asset_urls=True),
            message="Configuration published successfully"
        )
    except ValueError as e:
        return error_response(str(e))


@page_config_bp.route('/admin/<page_type>/history', methods=['GET'])
@jwt_required()
@role_required(UserRole.SUPER_ADMIN)
@feature_required('admin.page_configuration')
def get_version_history(page_type):
    """Get version history for a page type."""
    limit = request.args.get('limit', 10, type=int)
    versions = PageConfigService.get_version_history(page_type, limit)
    return success_response([v.to_dict() for v in versions])


@page_config_bp.route('/admin/<page_type>/restore/<version_id>', methods=['POST'])
@jwt_required()
@role_required(UserRole.SUPER_ADMIN)
@feature_required('admin.page_configuration')
def restore_version(page_type, version_id):
    """
    Restore a specific version to a new draft.
    Copies the selected version's configuration into the current draft.
    """
    try:
        user_id = str(current_user.id) if current_user else None
        draft = PageConfigService.restore_version(page_type, version_id, user_id)
        return success_response(
            draft.to_dict(include_asset_urls=True),
            message=f"Version restored to draft successfully"
        )
    except ValueError as e:
        return error_response(str(e))


@page_config_bp.route('/admin/audit-logs', methods=['GET'])
@jwt_required()
@role_required(UserRole.SUPER_ADMIN)
@feature_required('admin.page_configuration')
def get_audit_logs():
    """Get audit logs for all page configs."""
    page_type = request.args.get('page_type')
    limit = request.args.get('limit', 50, type=int)
    logs = PageConfigService.get_audit_logs(page_type, limit)
    return success_response([log.to_dict() for log in logs])


# ============== ASSET ENDPOINTS ==============

@page_config_bp.route('/admin/assets', methods=['GET'])
@jwt_required()
@role_required(UserRole.SUPER_ADMIN)
@feature_required('admin.page_configuration')
def list_assets():
    """List all assets, optionally filtered by type."""
    asset_type = request.args.get('type')
    assets = AssetService.get_all_assets(asset_type)
    return success_response([a.to_dict(include_url=True) for a in assets])


@page_config_bp.route('/admin/assets', methods=['POST'])
@jwt_required()
@role_required(UserRole.SUPER_ADMIN)
@feature_required('admin.page_configuration')
def upload_asset():
    """
    Upload a new asset to S3.
    
    Form data:
        - file: The file to upload
        - asset_type: Type of asset (logo, favicon, background_image, terms_document, privacy_document)
        - name: Display name for the asset
    """
    if 'file' not in request.files:
        return error_response("No file provided")
    
    file = request.files['file']
    if file.filename == '':
        return error_response("No file selected")
    
    asset_type = request.form.get('asset_type')
    name = request.form.get('name', file.filename)
    
    if not asset_type:
        return error_response("asset_type is required")
    
    try:
        # Upload to S3
        s3_data = S3Service.upload_file(file, asset_type, file.filename)
        
        # Create asset record
        user_id = str(current_user.id) if current_user else None
        asset = AssetService.create_asset(
            asset_type=asset_type,
            name=name,
            s3_data=s3_data,
            original_filename=file.filename,
            user_id=user_id
        )
        
        return created_response(
            asset.to_dict(include_url=True),
            message="Asset uploaded successfully",
        )
    except ValueError as e:
        return error_response(str(e))
    except Exception as e:
        return error_response(f"Upload failed: {str(e)}", status_code=500)


@page_config_bp.route('/admin/assets/<asset_id>', methods=['GET'])
@jwt_required()
@role_required(UserRole.SUPER_ADMIN)
@feature_required('admin.page_configuration')
def get_asset(asset_id):
    """Get asset by ID with presigned URL."""
    asset = AssetService.get_asset_by_id(asset_id)
    if not asset:
        return not_found_response("Asset")
    return success_response(asset.to_dict(include_url=True))


@page_config_bp.route('/admin/assets/<asset_id>', methods=['DELETE'])
@jwt_required()
@role_required(UserRole.SUPER_ADMIN)
@feature_required('admin.page_configuration')
def delete_asset(asset_id):
    """
    Delete an asset.
    Cannot delete assets used by LIVE configs.
    """
    if AssetService.is_asset_in_use(asset_id):
        return error_response(
            "Cannot delete asset that is used by a LIVE configuration",
            status_code=400,
        )

    user_id = str(current_user.id) if current_user else None
    if AssetService.delete_asset(asset_id, user_id):
        return success_response(None, message="Asset deleted successfully")

    return not_found_response("Asset")


@page_config_bp.route('/admin/assets/<asset_id>/url', methods=['GET'])
@jwt_required()
@role_required(UserRole.SUPER_ADMIN)
@feature_required('admin.page_configuration')
def get_asset_url(asset_id):
    """Get presigned URL for an asset."""
    expiration = request.args.get('expiration', 3600, type=int)
    asset = AssetService.get_asset_by_id(asset_id)
    if not asset:
        return not_found_response("Asset")

    url = asset.get_presigned_url(expiration)
    return success_response({'url': url, 'expires_in': expiration})


@page_config_bp.route('/admin/assets/<asset_id>/download', methods=['GET'])
@jwt_required()
@role_required(UserRole.SUPER_ADMIN)
@feature_required('admin.page_configuration')
def download_asset(asset_id):
    """
    Proxy-download an asset from S3.
    Bypasses CORS so the frontend can fetch file contents (e.g. for mammoth Word→HTML).
    """
    asset = AssetService.get_asset_by_id(asset_id)
    if not asset:
        return not_found_response("Asset")

    try:
        # Bucket-aware: a bare get_client() always returns the AWS client and
        # fails on MinIO-backed buckets.
        s3 = S3Service.get_client(asset.s3_bucket)
        s3_obj = s3.get_object(Bucket=asset.s3_bucket, Key=asset.s3_key)
        file_bytes = s3_obj['Body'].read()
        content_type = s3_obj.get('ContentType', asset.content_type or 'application/octet-stream')

        return Response(
            file_bytes,
            status=200,
            content_type=content_type,
            headers={
                'Content-Disposition': f'inline; filename="{asset.original_filename or "document"}"',
            }
        )
    except Exception as e:
        return error_response(f"Failed to download asset: {str(e)}", status_code=500)


@page_config_bp.route('/admin/assets/save-html-as-docx', methods=['POST'])
@jwt_required()
@role_required(UserRole.SUPER_ADMIN)
@feature_required('admin.page_configuration')
def save_html_as_docx():
    """
    Convert HTML content to a real .docx file, upload to S3, and create asset record.
    Used by the admin document editor (Quill) to save edited documents.

    Request JSON:
        - html_content: The HTML string from the editor
        - asset_type: 'terms_document' or 'privacy_document'
        - name: Display name for the asset
    """
    import io
    from docx import Document
    from docx.shared import Pt, Inches
    from bs4 import BeautifulSoup

    data = request.get_json()
    if not data or 'html_content' not in data:
        return error_response("html_content is required")

    html_content = data['html_content']
    asset_type = data.get('asset_type', 'terms_document')
    name = data.get('name', 'document')

    try:
        # Parse HTML and convert to docx
        soup = BeautifulSoup(html_content, 'html.parser')
        doc = Document()

        for element in soup.children:
            if element.name in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
                level = int(element.name[1]) - 1  # h1=0, h2=1, etc.
                doc.add_heading(element.get_text(), level=level)
            elif element.name == 'p':
                para = doc.add_paragraph()
                _add_inline_formatting(para, element)
            elif element.name in ('ul', 'ol'):
                for li in element.find_all('li', recursive=False):
                    para = doc.add_paragraph(style='List Bullet' if element.name == 'ul' else 'List Number')
                    _add_inline_formatting(para, li)
            elif element.name == 'br':
                doc.add_paragraph()
            elif isinstance(element, str) and element.strip():
                doc.add_paragraph(element.strip())

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Upload to S3
        filename = f"{name.replace(' ', '_')}.docx"
        buffer.name = filename
        buffer.content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

        s3_data = S3Service.upload_file(buffer, asset_type, filename)

        # Create asset record
        user_id = str(current_user.id) if current_user else None
        asset = AssetService.create_asset(
            asset_type=asset_type,
            name=name,
            s3_data=s3_data,
            original_filename=filename,
            user_id=user_id
        )

        return created_response(
            asset.to_dict(include_url=True),
            message="Document saved successfully",
        )
    except Exception as e:
        return error_response(f"Failed to save document: {str(e)}", status_code=500)


def _add_inline_formatting(para, element):
    """Helper: walk inline children of an HTML element and add runs with formatting."""
    from docx.shared import Pt
    from bs4 import NavigableString, Tag

    for child in element.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text.strip():
                para.add_run(text)
        elif isinstance(child, Tag):
            run = para.add_run(child.get_text())
            if child.name in ('strong', 'b'):
                run.bold = True
            elif child.name in ('em', 'i'):
                run.italic = True
            elif child.name == 'u':
                run.underline = True
            elif child.name == 's':
                run.font.strike = True
            elif child.name == 'a':
                run.underline = True  # Links shown as underlined text
